"""Expected values from a document the business hands over: CSV/TSV, XLSX (openpyxl), DOCX tables (python-docx),
Markdown tables. Output is an AgentTable plus an inferred grain (key column, dimensions, metric columns)."""
from __future__ import annotations
from ..textio import read_text
from ..install import install_cmd
import csv
import io
import os
import re
from typing import Any

from ..model import AgentTable, _coerce

_KEY_HEADERS = {"key", "issue_key", "issue", "ticket", "id", "jira_key"}
_KEY_VALUE = re.compile(r"^[A-Z][A-Z0-9_]+-\d+$")
_DATE_LIKE = re.compile(r"^\d{4}-\d{2}")
_DIM_NAMES = {"sprint", "status", "assignee", "team", "month", "week", "year", "quarter", "page", "visual", "category", "region"}


class ExpectError(Exception):
    def __init__(self, msg: str, hint: str = ""):
        super().__init__(msg)
        self.hint = hint


def normalize_header(h: Any) -> str:
    s = re.sub(r"[^0-9a-zA-Z]+", "_", str(h or "").strip()).strip("_").lower()
    return s or "col"


def _num(v: Any):
    if v is None:
        return None
    if isinstance(v, (int, float)) and not isinstance(v, bool):
        return v
    s = str(v).strip().replace(",", "")
    if s == "":
        return None
    pct = s.endswith("%")
    s = s.rstrip("%").strip()
    try:
        n = float(s)
    except ValueError:
        return None
    return int(n) if n.is_integer() and not pct and "." not in s else n


def _table(headers: list, rows: list[list], source: str, name: str = "expected") -> AgentTable:
    cols = [normalize_header(h) for h in headers]
    seen: dict[str, int] = {}
    for i, c in enumerate(cols):  # de-duplicate
        if c in seen:
            seen[c] += 1
            cols[i] = f"{c}_{seen[c]}"
        else:
            seen[c] = 0
    out_rows = []
    for r in rows:
        r = list(r) + [None] * (len(cols) - len(r))
        if all(v in (None, "") for v in r):
            continue
        out_rows.append([_coerce(v) if isinstance(v, str) else v for v in r[:len(cols)]])
    return AgentTable(name, cols, out_rows, source=source)


def load_expected(path: str, sheet: str | None = None, table_index: int = 0, name: str = "expected") -> AgentTable:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".csv", ".tsv", ".txt"):
        with open(path, encoding="utf-8-sig", newline="") as f:
            text = f.read()
        try:
            dialect = csv.Sniffer().sniff(text[:4096], delimiters=",\t;|")
            delim = dialect.delimiter
        except csv.Error:
            delim = "\t" if ext == ".tsv" else ","
        rows = list(csv.reader(io.StringIO(text), delimiter=delim))
        if not rows:
            raise ExpectError(f"{path} is empty")
        return _table(rows[0], rows[1:], path, name)
    if ext in (".xlsx", ".xlsm"):
        try:
            import openpyxl  # optional
        except ImportError:
            raise ExpectError("openpyxl is not installed", hint=install_cmd("uat")) from None
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb[sheet] if sheet else wb.active
        rows = [list(r) for r in ws.iter_rows(values_only=True)]
        rows = [[(v.isoformat() if hasattr(v, "isoformat") else v) for v in r] for r in rows]
        start = next((i for i, r in enumerate(rows) if sum(1 for v in r if v not in (None, "")) >= 2), None)
        if start is None:
            raise ExpectError(f"no header row found in {path}")
        return _table(rows[start], rows[start + 1:], f"{path}#{ws.title}", name)
    if ext == ".docx":
        try:
            import docx  # optional (python-docx)
        except ImportError:
            raise ExpectError("python-docx is not installed", hint=install_cmd("uat")) from None
        d = docx.Document(path)
        if table_index >= len(d.tables):
            raise ExpectError(f"{path} has {len(d.tables)} tables; --table-index {table_index} is out of range")
        tbl = d.tables[table_index]
        rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
        return _table(rows[0], rows[1:], f"{path}#table{table_index}", name)
    if ext in (".md", ".markdown"):
        blocks = markdown_tables(read_text(path))
        if table_index >= len(blocks):
            raise ExpectError(f"{path} has {len(blocks)} markdown tables; --table-index {table_index} is out of range")
        headers, rows = blocks[table_index]
        return _table(headers, rows, f"{path}#table{table_index}", name)
    raise ExpectError(f"unsupported file type {ext}", hint="use csv/tsv/xlsx/docx/md")


def markdown_tables(text: str) -> list[tuple[list[str], list[list[str]]]]:
    out, cur = [], []
    for line in text.splitlines() + [""]:
        if line.strip().startswith("|"):
            cur.append(line)
            continue
        if cur:
            rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in cur]
            rows = [r for r in rows if not all(re.match(r"^:?-{2,}:?$", c) for c in r if c)]
            if rows:
                out.append((rows[0], rows[1:]))
            cur = []
    return out


def infer_grain(t: AgentTable) -> dict:
    key = None
    for c in t.columns:
        if c in _KEY_HEADERS:
            key = c
            break
    if key is None:
        for i, c in enumerate(t.columns):
            vals = [r[i] for r in t.rows if r[i] not in (None, "")]
            if vals and all(isinstance(v, str) and _KEY_VALUE.match(v) for v in vals):
                key = c
                break
    metrics, dims = [], []
    for i, c in enumerate(t.columns):
        if c == key:
            continue
        vals = [r[i] for r in t.rows if r[i] not in (None, "")]
        nums = [v for v in (_num(v) for v in vals) if v is not None]
        if vals and len(nums) >= 0.8 * len(vals) and c not in _DIM_NAMES:
            metrics.append(c)
        elif vals and (c in _DIM_NAMES or all(isinstance(v, str) and _DATE_LIKE.match(v) for v in vals) or len(set(map(str, vals))) <= max(1, len(vals) // 2)):
            dims.append(c)
    g = {"key": key, "dims": dims, "metrics": metrics}
    if key is None:
        g["key"] = dims[0] if dims else (t.columns[0] if t.columns else None)
        g["warning"] = "no key column detected; using the first dimension"
    return g


def coerce_metrics(t: AgentTable, metrics: list[str]) -> None:
    for c in metrics:
        i = t.columns.index(c)
        for r in t.rows:
            r[i] = _num(r[i])
